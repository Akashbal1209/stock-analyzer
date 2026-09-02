import os
from datetime import datetime

# Import configuration
try:
    from config import REPORTS_FOLDER
except ImportError:
    REPORTS_FOLDER = "analysis_reports"


class ReportExporter:
    """Export formatted analysis reports to various file formats"""
    
    def __init__(self, reports_folder=None):
        self.reports_folder = reports_folder if reports_folder else REPORTS_FOLDER
        self._ensure_reports_folder()
    
    def _ensure_reports_folder(self):
        """Create reports folder if it doesn't exist"""
        if not os.path.exists(self.reports_folder):
            os.makedirs(self.reports_folder)
            print(f"✅ Created folder: {self.reports_folder}")
    
    def save_report(self, symbol, report_text, format='txt'):
        """
        Save report in specified format
        
        Parameters:
        - symbol: Stock symbol
        - report_text: The formatted analysis text
        - format: 'txt', 'html', or 'docx'
        
        Returns:
        - Path to saved file
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        if format == 'txt':
            return self._save_txt(symbol, report_text, timestamp)
        elif format == 'html':
            return self._save_html(symbol, report_text, timestamp)
        elif format == 'docx':
            return self._save_docx(symbol, report_text, timestamp)
        else:
            print(f"❌ Unknown format: {format}")
            return None
    
    def _save_txt(self, symbol, report_text, timestamp):
        """Save as plain text file"""
        filename = f"{self.reports_folder}/{symbol}_Analysis_{timestamp}.txt"
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(report_text)
            
            full_path = os.path.abspath(filename)
            print(f"\n✅ TXT Report saved: {full_path}")
            return full_path
            
        except Exception as e:
            print(f"❌ Error saving TXT: {e}")
            return None
    
    def _save_html(self, symbol, report_text, timestamp):
        """Save as HTML file with nice formatting"""
        filename = f"{self.reports_folder}/{symbol}_Analysis_{timestamp}.html"
        
        html_content = self._text_to_html(symbol, report_text)
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            full_path = os.path.abspath(filename)
            print(f"\n✅ HTML Report saved: {full_path}")
            print(f"   💡 Open this file in your browser to view")
            return full_path
            
        except Exception as e:
            print(f"❌ Error saving HTML: {e}")
            return None
    
    def _text_to_html(self, symbol, text):
        """Convert plain text report to beautiful HTML"""
        html_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        html_text = html_text.replace('\n', '<br>\n')
        
        # Style the special characters
        html_text = html_text.replace('✅', '<span style="color: #22c55e;">✅</span>')
        html_text = html_text.replace('❌', '<span style="color: #ef4444;">❌</span>')
        html_text = html_text.replace('⚠️', '<span style="color: #f59e0b;">⚠️</span>')
        html_text = html_text.replace('🟢', '<span style="color: #22c55e;">🟢</span>')
        html_text = html_text.replace('🔴', '<span style="color: #ef4444;">🔴</span>')
        html_text = html_text.replace('🟡', '<span style="color: #f59e0b;">🟡</span>')
        
        html_template = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{symbol} - Stock Analysis Report</title>
    <style>
        body {{
            font-family: 'Courier New', monospace;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            margin: 0;
            padding: 20px;
            min-height: 100vh;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
        }}
        .header {{
            text-align: center;
            color: #667eea;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 3px solid #667eea;
        }}
        .report {{
            white-space: pre-wrap;
            line-height: 1.4;
            font-size: 14px;
            color: #1f2937;
        }}
        .print-btn {{
            position: fixed;
            bottom: 30px;
            right: 30px;
            background: #667eea;
            color: white;
            padding: 15px 30px;
            border: none;
            border-radius: 50px;
            cursor: pointer;
            font-size: 16px;
            font-weight: bold;
            box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
            transition: all 0.3s;
        }}
        .print-btn:hover {{
            background: #5568d3;
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(102, 126, 234, 0.6);
        }}
        @media print {{
            body {{
                background: white;
                padding: 0;
            }}
            .container {{
                box-shadow: none;
                padding: 20px;
            }}
            .print-btn {{
                display: none;
            }}
        }}
    </style>
</head>
<body>
    <button class="print-btn" onclick="window.print()">🖨️ Print Report</button>
    <div class="container">
        <div class="header">
            <h1>📊 {symbol} - STOCK ANALYSIS REPORT</h1>
            <p>Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</p>
        </div>
        <div class="report">{html_text}</div>
    </div>
</body>
</html>"""
        return html_template
    
    def _save_docx(self, symbol, report_text, timestamp):
        """Save as Word document (requires python-docx)"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("\n⚠️ python-docx not installed!")
            print("   Install it with: pip install python-docx")
            print("   For now, saving as TXT instead...")
            return self._save_txt(symbol, report_text, timestamp)
        
        filename = f"{self.reports_folder}/{symbol}_Analysis_{timestamp}.docx"
        
        try:
            doc = Document()
            
            title = doc.add_heading(f'{symbol} - STOCK ANALYSIS REPORT', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            lines = report_text.split('\n')
            for line in lines:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                
                if '✅' in line:
                    run.font.color.rgb = RGBColor(34, 197, 94)
                elif '❌' in line:
                    run.font.color.rgb = RGBColor(239, 68, 68)
                elif '⚠️' in line:
                    run.font.color.rgb = RGBColor(245, 158, 11)
            
            doc.save(filename)
            
            full_path = os.path.abspath(filename)
            print(f"\n✅ DOCX Report saved: {full_path}")
            print(f"   💡 Open this file in Microsoft Word")
            return full_path
            
        except Exception as e:
            print(f"❌ Error saving DOCX: {e}")
            return None
    
    def save_all_formats(self, symbol, report_text):
        """Save report in all available formats"""
        print("\n" + "="*60)
        print("   📁 SAVING REPORTS IN MULTIPLE FORMATS")
        print("="*60)
        
        paths = {}
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        paths['txt'] = self._save_txt(symbol, report_text, timestamp)
        paths['html'] = self._save_html(symbol, report_text, timestamp)
        paths['docx'] = self._save_docx(symbol, report_text, timestamp)
        
        return paths
    
    def auto_open_file(self, filepath):
        """Auto-open the saved file in default program"""
        try:
            import platform
            import subprocess
            
            system = platform.system()
            
            if system == 'Windows':
                os.startfile(filepath)
                print(f"   📂 Opened file automatically!")
            elif system == 'Darwin':
                subprocess.run(['open', filepath])
                print(f"   📂 Opened file automatically!")
            elif system == 'Linux':
                subprocess.run(['xdg-open', filepath])
                print(f"   📂 Opened file automatically!")
                
        except Exception as e:
            print(f"   💡 Could not auto-open file. Please open manually.")


# Test module
if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("   TESTING REPORT EXPORTER")
    print("=" * 60)
    
    exporter = ReportExporter()
    
    sample_report = """
═══════════════════════════════════════════════════════════
   📊 RELIANCE - COMPLETE STOCK ANALYSIS
═══════════════════════════════════════════════════════════

   💰 Current Price: ₹2450.50
   📅 Analysis Date: 2024-01-15

   🎯 OVERALL SIGNAL: 🟢 BUY
   📊 Overall Score: 72.5/100

═══════════════════════════════════════════════════════════
   📋 FINAL SUMMARY
═══════════════════════════════════════════════════════════

   ┌───────────────────────────────────────────────────────┐
   │ FUNDAMENTAL : ✅ Strong (75/100)                      │
   │ HISTORICAL  : ⚠️ Neutral (55/100)                    │
   │ TECHNICAL   : ✅ Bullish (68/100)                    │
   └───────────────────────────────────────────────────────┘
"""
    
    paths = exporter.save_all_formats('RELIANCE', sample_report)
    
    print("\n" + "="*60)
    print("   ✅ TEST COMPLETED!")
    print("="*60)
    print(f"\nSaved files in folder: {exporter.reports_folder}")
